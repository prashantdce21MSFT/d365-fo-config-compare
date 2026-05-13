"""Convert the Team Guide markdown to a formatted Word document."""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

MD_PATH = r"C:\D365 Configuration Drift Analysis\TEAM_GUIDE_D365_Config_Drift_Analysis.md"
DOCX_PATH = r"C:\D365 Configuration Drift Analysis\TEAM_GUIDE_D365_Config_Drift_Analysis.docx"


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shd)


def add_table_from_rows(doc, header_row, data_rows):
    """Add a formatted table to the document."""
    cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for i, h in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = h.strip()
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, "1F4E79")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for ri, row in enumerate(data_rows):
        for ci in range(cols):
            cell = table.rows[ri + 1].cells[ci]
            val = row[ci].strip() if ci < len(row) else ""
            # Handle inline code
            if '`' in val:
                p = cell.paragraphs[0]
                parts = re.split(r'(`[^`]+`)', val)
                for part in parts:
                    if part.startswith('`') and part.endswith('`'):
                        run = p.add_run(part[1:-1])
                        run.font.name = 'Consolas'
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
                    else:
                        run = p.add_run(part)
                        run.font.size = Pt(9)
            else:
                cell.text = val
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
            if ri % 2 == 1:
                set_cell_shading(cell, "F2F2F2")

    return table


def process_inline(paragraph, text):
    """Process inline markdown: **bold**, `code`, normal text."""
    # Split on bold and code patterns
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
        else:
            paragraph.add_run(part)


def convert():
    with open(MD_PATH, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(2)

    # Title page
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_heading('D365 Configuration Drift Analysis', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading('Team Guide', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Using GitHub Copilot (Claude Code) + D365 MCP\nfor Automated Form Comparison').italic = True
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Author: ').bold = True
    p.add_run('Prashant Verma, QA - Acme Asia SmartCore')
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Last Updated: ').bold = True
    p.add_run('2026-04-30')
    doc.add_page_break()

    i = 0
    in_code_block = False
    code_lines = []
    code_lang = ""
    in_table = False
    table_header = []
    table_rows = []

    # Skip the first few lines (title, subtitle, author, date, ---)
    # Find where actual content starts (after the ToC separator)
    start = 0
    hr_count = 0
    for idx, line in enumerate(lines):
        if line.strip() == '---':
            hr_count += 1
            if hr_count >= 3:  # After title block + first separator + ToC separator
                start = idx + 1
                break
    if start == 0:
        start = 0

    i = start
    while i < len(lines):
        line = lines[i]
        raw = line.rstrip('\n')
        stripped = raw.strip()

        # Code block toggle
        if stripped.startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                # Add light gray background via paragraph shading
                pPr = p._element.get_or_add_pPr()
                shd = pPr.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'F5F5F5',
                    qn('w:val'): 'clear',
                })
                pPr.append(shd)
                p.paragraph_format.left_indent = Inches(0.3)
                code_lines = []
                in_code_block = False
            else:
                # Flush any pending table
                if in_table and table_header:
                    add_table_from_rows(doc, table_header, table_rows)
                    doc.add_paragraph()
                    in_table = False
                    table_header = []
                    table_rows = []
                in_code_block = True
                code_lang = stripped[3:].strip()
            i += 1
            continue

        if in_code_block:
            code_lines.append(raw)
            i += 1
            continue

        # Table detection
        if '|' in stripped and stripped.startswith('|') and stripped.endswith('|'):
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            # Check if separator row
            if all(re.match(r'^[-:]+$', c) for c in cells):
                i += 1
                continue
            if not in_table:
                in_table = True
                table_header = cells
                table_rows = []
            else:
                table_rows.append(cells)
            i += 1
            continue
        else:
            # Flush pending table
            if in_table and table_header:
                add_table_from_rows(doc, table_header, table_rows)
                doc.add_paragraph()
                in_table = False
                table_header = []
                table_rows = []

        # Empty line
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if stripped == '---':
            doc.add_paragraph('_' * 60)
            i += 1
            continue

        # Headings
        if stripped.startswith('#'):
            m = re.match(r'^(#{1,4})\s+(.+)$', stripped)
            if m:
                level = len(m.group(1))
                text = m.group(2)
                # Remove markdown links from heading
                text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
                if level == 1:
                    doc.add_page_break()
                doc.add_heading(text, level=min(level, 4))
                i += 1
                continue

        # Checkbox list items
        if stripped.startswith('- [ ]') or stripped.startswith('- [x]'):
            checked = stripped.startswith('- [x]')
            text = stripped[5:].strip()
            p = doc.add_paragraph(style='List Bullet')
            marker = "[x] " if checked else "[ ] "
            run = p.add_run(marker)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            process_inline(p, text)
            i += 1
            continue

        # Bullet list
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            process_inline(p, text)
            i += 1
            continue

        # Numbered list
        m = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if m:
            text = m.group(2)
            p = doc.add_paragraph(style='List Number')
            process_inline(p, text)
            i += 1
            continue

        # Blockquote
        if stripped.startswith('>'):
            text = stripped.lstrip('> ')
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        process_inline(p, stripped)
        i += 1

    # Flush final table
    if in_table and table_header:
        add_table_from_rows(doc, table_header, table_rows)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    doc.save(DOCX_PATH)
    print(f"Saved: {DOCX_PATH}")


if __name__ == '__main__':
    convert()
